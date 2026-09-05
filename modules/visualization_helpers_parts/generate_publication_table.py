"""Create publication-grade descriptive tables from mixed variable specifications."""

from __future__ import annotations

from collections import OrderedDict
from pathlib import Path
import textwrap

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

# python-docx is optional — only required when "docx" is in save_formats
try:
    from docx import Document as _Document
    from docx.shared import Pt as _Pt, Cm as _Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN
    from docx.enum.section import WD_ORIENT as _WD_ORIENT
    from docx.oxml.ns import qn as _qn, nsdecls as _nsdecls
    from docx.oxml import parse_xml as _parse_xml
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False


def _is_nan_section_marker(value) -> bool:
    """Return True when `value` is a NaN-like section separator marker."""
    return isinstance(value, (float, np.floating)) and pd.isna(value)


def _resolve_columns(dataframe: pd.DataFrame, columns, include_total: bool) -> OrderedDict:
    """Resolve user-provided column masks to aligned boolean Series."""
    resolved = OrderedDict()

    if columns is None:
        resolved["Total"] = pd.Series(True, index=dataframe.index)
        return resolved

    for label, mask in columns.items():
        if callable(mask):
            mask = mask(dataframe)
        if isinstance(mask, pd.Series):
            series_mask = mask.reindex(dataframe.index, fill_value=False)
        else:
            series_mask = pd.Series(mask, index=dataframe.index)
        resolved[label] = series_mask.fillna(False).astype(bool)

    if include_total and "Total" not in resolved:
        resolved["Total"] = pd.Series(True, index=dataframe.index)

    return resolved


def _format_continuous(values: pd.Series, digits: int) -> str:
    """Format continuous values as mean +/- SD."""
    if len(values) == 0:
        return "--"
    mean_value = values.mean()
    std_value = values.std()
    return f"{mean_value:.{digits}f} +/- {std_value:.{digits}f}"


def _format_count_percentage(series: pd.Series, code, total_n: int, pct_digits: int) -> str:
    """Format categorical rows as n (pct%)."""
    if total_n == 0:
        return "--"
    count = int((series == code).sum())
    pct = (count / total_n) * 100
    return f"{count} ({pct:.{pct_digits}f}%)"


def _build_table_rows(
    dataframe: pd.DataFrame,
    resolved_columns: OrderedDict,
    sections: OrderedDict,
    continuous_digits: int,
    pct_digits: int,
) -> pd.DataFrame:
    """Build long-format table rows with metadata for rendering/export."""
    col_labels = list(resolved_columns.keys())
    rows = []

    for section, var_spec in sections.items():
        if _is_nan_section_marker(var_spec):
            row = {
                "section": section,
                "variable": "",
                "code": "",
                "row_type": "section_break",
                "Characteristic": section,
            }
            for col in col_labels:
                row[col] = ""
            rows.append(row)
            continue

        if isinstance(var_spec, str):
            section_row = {
                "section": section,
                "variable": var_spec,
                "code": "",
                "row_type": "section_header",
                "Characteristic": section,
            }
            for col in col_labels:
                section_row[col] = ""
            rows.append(section_row)

            value_row = {
                "section": section,
                "variable": var_spec,
                "code": "",
                "row_type": "continuous",
                "Characteristic": "",
            }
            for col, mask in resolved_columns.items():
                values = pd.to_numeric(dataframe.loc[mask, var_spec], errors="coerce").dropna()
                value_row[col] = _format_continuous(values, digits=continuous_digits)
            rows.append(value_row)
            continue

        if isinstance(var_spec, dict):
            if len(var_spec) != 1:
                raise ValueError(
                    f"Section '{section}' must map to exactly one categorical column, got {len(var_spec)}."
                )
            colname, mapping = next(iter(var_spec.items()))

            if len(mapping) == 1:
                code, _label = next(iter(mapping.items()))
                single_row = {
                    "section": section,
                    "variable": colname,
                    "code": code,
                    "row_type": "binary_indicator",
                    "Characteristic": section,
                }
                for col, mask in resolved_columns.items():
                    values = dataframe.loc[mask, colname]
                    total_n = int(mask.sum())
                    single_row[col] = _format_count_percentage(values, code, total_n, pct_digits)
                rows.append(single_row)
            else:
                section_row = {
                    "section": section,
                    "variable": colname,
                    "code": "",
                    "row_type": "section_header",
                    "Characteristic": section,
                }
                for col in col_labels:
                    section_row[col] = ""
                rows.append(section_row)

                for code, label in mapping.items():
                    level_row = {
                        "section": section,
                        "variable": colname,
                        "code": code,
                        "row_type": "categorical_level",
                        "Characteristic": str(label),
                    }
                    for col, mask in resolved_columns.items():
                        values = dataframe.loc[mask, colname]
                        total_n = int(mask.sum())
                        level_row[col] = _format_count_percentage(values, code, total_n, pct_digits)
                    rows.append(level_row)
            continue

        raise TypeError(
            f"Unsupported section spec for '{section}'. Expected str, dict, or np.nan marker."
        )

    return pd.DataFrame(rows)


def _analysis_columns_for_tests(resolved_columns: OrderedDict) -> list[str]:
    """Return comparison columns used for inferential tests (exclude Total)."""
    return [col for col in resolved_columns.keys() if str(col).strip().lower() != "total"]


def _compute_variable_test(
    dataframe: pd.DataFrame,
    resolved_columns: OrderedDict,
    variable: str,
    variable_kind: str,
    continuous_test: str = "anova",
) -> dict:
    """
    Compute inferential test for a variable across analysis groups.

    Returns dict with keys: statistic_type, test_statistic, p_value.
    continuous_test: "anova" (default) or "mannwhitneyu" (2-group only).
    """
    from scipy.stats import chi2_contingency, f_oneway, mannwhitneyu

    group_cols = _analysis_columns_for_tests(resolved_columns)
    if len(group_cols) < 2:
        return {"statistic_type": "", "test_statistic": np.nan, "p_value": np.nan}

    if variable_kind == "continuous":
        groups = []
        for col in group_cols:
            vals = pd.to_numeric(dataframe.loc[resolved_columns[col], variable], errors="coerce").dropna()
            if len(vals) > 0:
                groups.append(vals.values)

        if len(groups) < 2:
            return {"statistic_type": continuous_test, "test_statistic": np.nan, "p_value": np.nan}

        if continuous_test == "mannwhitneyu" and len(groups) == 2:
            stat, p_val = mannwhitneyu(groups[0], groups[1], alternative="two-sided")
            return {"statistic_type": "mannwhitneyu", "test_statistic": float(stat), "p_value": float(p_val)}

        stat, p_val = f_oneway(*groups)
        return {"statistic_type": "anova", "test_statistic": float(stat), "p_value": float(p_val)}

    if variable_kind == "categorical":
        grouped_series = []
        category_index = pd.Index([])
        for col in group_cols:
            vals = dataframe.loc[resolved_columns[col], variable].dropna()
            grouped_series.append(vals)
            category_index = category_index.union(pd.Index(vals.unique()))

        if len(category_index) < 2:
            return {"statistic_type": "chi_square", "test_statistic": np.nan, "p_value": np.nan}

        contingency_rows = []
        for vals in grouped_series:
            counts = vals.value_counts().reindex(category_index, fill_value=0)
            contingency_rows.append(counts.values)

        contingency = np.asarray(contingency_rows)
        if contingency.shape[0] < 2 or contingency.shape[1] < 2:
            return {"statistic_type": "chi_square", "test_statistic": np.nan, "p_value": np.nan}

        stat, p_val, _dof, _expected = chi2_contingency(contingency)
        return {"statistic_type": "chi_square", "test_statistic": float(stat), "p_value": float(p_val)}

    return {"statistic_type": "", "test_statistic": np.nan, "p_value": np.nan}


def _compute_binary_level_test(
    dataframe: pd.DataFrame,
    resolved_columns: OrderedDict,
    variable: str,
    code,
) -> dict:
    """Chi-square (or Fisher exact) test for membership in a single category vs group.

    When the contingency table is exactly 2×2 (two
    comparison groups), Fisher's exact test replaces chi-square whenever the
    expected-count assumption is violated (any expected cell < 5).  For tables
    with more than two groups, chi-square is always used and the usual 'NA'
    is returned on assumption violation.
    """
    from scipy.stats import chi2_contingency, fisher_exact

    group_cols = _analysis_columns_for_tests(resolved_columns)
    if len(group_cols) < 2:
        return {
            "statistic_type": "chi_square", "test_statistic": np.nan,
            "p_value": np.nan, "assumption_violated": False,
        }

    contingency_rows = []
    for col in group_cols:
        vals = dataframe.loc[resolved_columns[col], variable]
        yes = (vals == code).sum()
        no = vals.notna().sum() - yes
        contingency_rows.append([no, yes])

    contingency = np.asarray(contingency_rows)
    if contingency.shape[0] < 2 or contingency.shape[1] < 2:
        return {
            "statistic_type": "chi_square", "test_statistic": np.nan,
            "p_value": np.nan, "assumption_violated": False,
        }
    if (
        contingency.sum() == 0
        or contingency[:, 1].sum() == 0
        or contingency[:, 0].sum() == 0
        or (contingency.sum(axis=1) == 0).any()
    ):
        return {
            "statistic_type": "chi_square", "test_statistic": np.nan,
            "p_value": np.nan, "assumption_violated": False,
        }

    stat, p_val, _dof, expected = chi2_contingency(contingency)
    if (expected < 5).any():
        if contingency.shape == (2, 2):
            # Fisher's exact test is appropriate for 2×2 tables when chi-square
            # expected-count assumptions are violated.
            _, fe_p = fisher_exact(contingency)
            return {
                "statistic_type": "fisher_exact", "test_statistic": np.nan,
                "p_value": float(fe_p), "assumption_violated": False,
            }
        return {
            "statistic_type": "chi_square", "test_statistic": np.nan,
            "p_value": np.nan, "assumption_violated": True,
        }
    return {
        "statistic_type": "chi_square", "test_statistic": float(stat),
        "p_value": float(p_val), "assumption_violated": False,
    }


def _compute_omnibus_chi_square(
    dataframe: pd.DataFrame,
    resolved_columns: OrderedDict,
    variable: str,
    min_expected: float = 5.0,
) -> dict:
    """
    Omnibus chi-square (or Fisher exact) test across all levels of a categorical variable.

    Builds a (n_groups × n_levels) contingency table, drops any level (column)
    where any expected cell count is below `min_expected`, then runs chi-square
    on the remaining levels.  Returns assumption_violated=True if fewer than 2
    levels survive after dropping.

    When the *original* (un-dropped) contingency table
    is exactly 2×2, Fisher's exact test is run on the full table instead of
    returning 'NA'.  For tables with more than 2 levels, scipy provides no
    general R×C Fisher exact, so assumption_violated=True / NA is still returned.

    Returns dict with keys: statistic_type, test_statistic, p_value,
    assumption_violated, n_levels_dropped.
    """
    from scipy.stats import chi2_contingency, fisher_exact

    group_cols = _analysis_columns_for_tests(resolved_columns)
    if len(group_cols) < 2:
        return {
            "statistic_type": "chi_square", "test_statistic": np.nan,
            "p_value": np.nan, "assumption_violated": True, "n_levels_dropped": 0,
        }

    grouped_series = []
    category_index = pd.Index([])
    for col in group_cols:
        vals = dataframe.loc[resolved_columns[col], variable].dropna()
        grouped_series.append(vals)
        category_index = category_index.union(pd.Index(vals.unique()))

    if len(category_index) < 2:
        return {
            "statistic_type": "chi_square", "test_statistic": np.nan,
            "p_value": np.nan, "assumption_violated": True, "n_levels_dropped": 0,
        }

    contingency_rows = []
    for vals in grouped_series:
        counts = vals.value_counts().reindex(category_index, fill_value=0)
        contingency_rows.append(counts.values)
    contingency = np.asarray(contingency_rows)

    # Guard: if any group (row) has zero total observations across all levels,
    # chi2_contingency cannot compute expected values — return gracefully.
    if (contingency.sum(axis=1) == 0).any():
        return {
            "statistic_type": "chi_square", "test_statistic": np.nan,
            "p_value": np.nan, "assumption_violated": True, "n_levels_dropped": 0,
        }

    # Compute expected counts for the full table to identify violating levels
    _, _, _dof, expected = chi2_contingency(contingency)

    # Drop levels where any expected cell is below min_expected
    keep_mask = (expected >= min_expected).all(axis=0)
    n_levels_dropped = int((~keep_mask).sum())
    contingency_kept = contingency[:, keep_mask]

    if contingency_kept.shape[1] < 2:
        # Fewer than 2 levels survive the expected-count filter.
        # If the original table was exactly 2×2 and Fisher exact is requested,
        # run it on the full (unfiltered) table instead.
        if contingency.shape == (2, 2):
            _, fe_p = fisher_exact(contingency)
            return {
                "statistic_type": "fisher_exact", "test_statistic": np.nan,
                "p_value": float(fe_p), "assumption_violated": False,
                "n_levels_dropped": n_levels_dropped,
            }
        return {
            "statistic_type": "chi_square", "test_statistic": np.nan,
            "p_value": np.nan, "assumption_violated": True,
            "n_levels_dropped": n_levels_dropped,
        }

    stat, p_val, _dof2, _exp2 = chi2_contingency(contingency_kept)
    return {
        "statistic_type": "chi_square", "test_statistic": float(stat),
        "p_value": float(p_val), "assumption_violated": False,
        "n_levels_dropped": n_levels_dropped,
    }


def _format_p(val, assumption_violated: bool = False) -> str:
    """
    Format a p-value for display.

    Returns 'NA' when chi-square expected-count assumptions are violated,
    '' for NaN (test not applicable), '<0.001' or a 3-decimal string otherwise.
    """
    if assumption_violated:
        return "NA"
    if pd.isna(val):
        return ""
    if val < 0.001:
        return "<0.001"
    return f"{val:.3f}"


def _attach_row_level_test_stats(
    row_df: pd.DataFrame,
    dataframe: pd.DataFrame,
    resolved_columns: OrderedDict,
    continuous_test: str = "anova",
    include_omnibus_chi_square: bool = True,
) -> pd.DataFrame:
    """Add statistic_type / test_statistic / p_value / assumption_violated columns aligned to table rows.

    For multi-level categorical section_header rows, if include_omnibus_chi_square is True,
    an omnibus chi-square is run across all levels (dropping those that violate the expected-
    count assumption) and the result is stored on the section_header row.  Individual
    categorical_level rows show 'NA' when their per-cell expected counts are too low.

    Fisher's exact test replaces the 'NA' for any 2×2
    contingency table (binary indicators and omnibus tests on binary variables).
    """
    out = row_df.copy()
    out["statistic_type"] = ""
    out["test_statistic"] = np.nan
    out["p_value"] = np.nan
    out["chi_square_p_value"] = np.nan
    out["assumption_violated"] = False

    valid_rows = out["row_type"].isin({"continuous", "binary_indicator", "categorical_level"})
    for idx, row in out[valid_rows].iterrows():
        row_type = row["row_type"]
        variable = row["variable"]
        if not variable:
            continue
        if row_type == "continuous":
            stats = _compute_variable_test(dataframe, resolved_columns, variable, "continuous", continuous_test)
            out.at[idx, "assumption_violated"] = False
        else:
            stats = _compute_binary_level_test(dataframe, resolved_columns, variable, row["code"],
)
            out.at[idx, "assumption_violated"] = stats.get("assumption_violated", False)

        out.at[idx, "statistic_type"] = stats["statistic_type"]
        out.at[idx, "test_statistic"] = stats["test_statistic"]
        out.at[idx, "p_value"] = stats["p_value"]
        if stats["statistic_type"] in {"chi_square", "anova", "mannwhitneyu", "fisher_exact"}:
            out.at[idx, "chi_square_p_value"] = stats["p_value"]

    # Omnibus chi-square for section_header rows of multi-level categoricals only.
    # Continuous variables also produce a section_header row (used for the bold label),
    # but they are followed by a "continuous" child row, not "categorical_level" rows.
    # We distinguish them by checking whether the section has any categorical_level children.
    if include_omnibus_chi_square:
        for idx, row in out[out["row_type"] == "section_header"].iterrows():
            variable = row["variable"]
            if not variable:
                continue
            section_name = row["section"]
            has_cat_levels = (
                (out["section"] == section_name) & (out["row_type"] == "categorical_level")
            ).any()
            if not has_cat_levels:
                continue  # skip section headers for continuous variables
            omnibus = _compute_omnibus_chi_square(dataframe, resolved_columns, variable,
)
            out.at[idx, "statistic_type"] = omnibus["statistic_type"]
            out.at[idx, "test_statistic"] = omnibus["test_statistic"]
            out.at[idx, "p_value"] = omnibus["p_value"]
            out.at[idx, "chi_square_p_value"] = omnibus["p_value"]
            out.at[idx, "assumption_violated"] = omnibus["assumption_violated"]

    return out


def _build_styler(display_df: pd.DataFrame, row_meta: pd.DataFrame, table_caption: str):
    """Create a styled HTML table approximating LaTeX manuscript style."""

    def style_row(row: pd.Series):
        row_type = row_meta.loc[row.name, "row_type"]
        if row_type in {"section_header", "section_break"}:
            return [
                "font-style: italic; font-weight: 700; background-color: #f4f4f4; border-top: 1px solid #555;"
            ] * len(row)
        if row_type == "binary_indicator":
            return ["font-weight: 600;"] + [""] * (len(row) - 1)
        return [""] * len(row)

    table_styles = [
        {
            "selector": "table",
            "props": [
                ("border-collapse", "collapse"),
                ("border-top", "2px solid #222"),
                ("border-bottom", "2px solid #222"),
                ("width", "100%"),
            ],
        },
        {
            "selector": "th",
            "props": [
                ("font-family", "Times New Roman, serif"),
                ("font-size", "12pt"),
                ("font-weight", "700"),
                ("border-bottom", "1.5px solid #333"),
                ("padding", "6px 8px"),
                ("text-align", "center"),
                ("border-left", "none"),
                ("border-right", "none"),
            ],
        },
        {
            "selector": "td",
            "props": [
                ("font-family", "Times New Roman, serif"),
                ("font-size", "11pt"),
                ("padding", "4px 8px"),
                ("border-bottom", "0.5px solid #8a8a8a"),
                ("border-left", "none"),
                ("border-right", "none"),
            ],
        },
        {"selector": "tbody td:first-child", "props": [("text-align", "left")]},
    ]
    if "P-value" in display_df.columns:
        col_idx = list(display_df.columns).index("P-value")
        table_styles.append(
            {
                "selector": f"th.col_heading.level0.col{col_idx}",
                "props": [("font-style", "italic")],
            }
        )

    styler = (
        display_df.style.hide(axis="index")
        .set_table_styles(table_styles)
        .apply(style_row, axis=1)
        .set_properties(subset=display_df.columns[1:], **{"text-align": "center"})
    )
    if table_caption:
        styler = styler.set_caption(table_caption)
    return styler


def _save_png_table(
    display_df: pd.DataFrame,
    row_meta: pd.DataFrame,
    savepath: Path,
    table_caption: str,
    wrap_width: int,
):
    """Save the styled table as a high-resolution PNG using matplotlib."""
    plt.rcParams["font.family"] = "Arial"
    wrapped = display_df.copy()
    wrapped["Characteristic"] = wrapped["Characteristic"].fillna("").map(
        lambda x: "\n".join(textwrap.wrap(str(x), width=wrap_width)) if str(x) else ""
    )

    n_rows = len(wrapped)
    n_cols = len(wrapped.columns)
    fig_width = max(3 + n_cols, 3.5 + (n_cols * 2.25))
    fig_height = max(3.8, 0.5 + (n_rows * 0.16))

    fig, ax = plt.subplots(figsize=(fig_width, fig_height))
    ax.axis("off")

    table = ax.table(
        cellText=wrapped.values,
        colLabels=wrapped.columns,
        loc="upper left",
        cellLoc="left",
        colLoc="center",
    )
    if "P-value" in wrapped.columns:
        col_idx = list(wrapped.columns).index("P-value")
        table[(0, col_idx)].get_text().set_fontstyle("italic")
    table.auto_set_font_size(False)
    fontsize = 10.5 if n_rows <= 35 else 9.5
    table.set_fontsize(fontsize)

    if n_cols <= 2:
        first_col_width = 0.86
    elif n_cols <= 4:
        first_col_width = 0.44
    else:
        first_col_width = 0.39
    remaining_width = 1.0 - first_col_width
    other_col_width = remaining_width / (n_cols - 1)

    # Allocate heights by row text density so wrapped rows do not overlap.
    row_weights = []
    for row_i in range(n_rows):
        line_count = max(1, wrapped.iloc[row_i, 0].count("\n") + 1)
        row_type = row_meta.iloc[row_i]["row_type"]
        weight = 1.0 + 0.78 * (line_count - 1)
        if row_type in {"section_header", "section_break"}:
            weight += 0.28
        row_weights.append(weight)

    header_weight = 1.35
    table_fill = 0.95
    per_weight = table_fill / (header_weight + sum(row_weights))
    header_height = header_weight * per_weight

    for (row_i, col_i), cell in table.get_celld().items():
        cell.set_facecolor("white")
        cell.set_edgecolor("#222222")
        cell.set_width(first_col_width if col_i == 0 else other_col_width)

        if row_i == 0:
            cell.visible_edges = "TB"
            cell.set_text_props(weight="bold", ha="center", va="center")
            cell.set_linewidth(1.5)
            cell.set_height(header_height)
            continue

        row_type = row_meta.iloc[row_i - 1]["row_type"]
        row_height = row_weights[row_i - 1] * per_weight
        cell.set_height(row_height)

        if col_i == 0:
            cell.set_text_props(ha="left")
        else:
            cell.set_text_props(ha="center")

        if row_type in {"section_header", "section_break"}:
            cell.visible_edges = "TB"
            if col_i == 0:
                cell.set_text_props(style="italic", weight="bold", ha="left")
            cell.set_linewidth(1.0)
        else:
            cell.visible_edges = "B"
            cell.set_linewidth(0.45)

    fig.subplots_adjust(left=0.02, right=0.98, top=0.95, bottom=0.02)
    fig.savefig(savepath, dpi=400, bbox_inches="tight")
    plt.close(fig)


def generate_publication_table(
    dataframe: pd.DataFrame,
    columns: dict | None = None,
    sections: OrderedDict | dict | None = None,
    table_caption: str = "",
    table_label: str = "tab:summary",
    savepath: str | None = None,
    save_formats=("png", "html", "tex", "csv"),
    include_total: bool = True,
    continuous_digits: int = 2,
    pct_digits: int = 1,
    wrap_width: int = 60,
    include_test_stats_csv: bool = True,
    continuous_test: str = "anova",
    include_omnibus_chi_square: bool = True,
):
    """
    Build a publication-grade descriptive table and optionally save render/export files.

    Args:
        dataframe: Source dataframe.
        columns: Dict of column-name -> boolean mask (or callable returning mask).
        sections: Ordered mapping of row labels to variable specs:
            - str: continuous variable (mean +/- SD)
            - dict: {'colname': {code: label, ...}} categorical row(s)
            - np.nan: section break row
        table_caption: Display caption.
        table_label: LaTeX label used in `.tex` output.
        savepath: Base path without extension (or with a known extension).
        save_formats: Iterable of output formats among {'png','html','tex','csv'}.
        include_total: Append "Total" column when not provided.
        continuous_digits: Decimal places for mean/SD.
        pct_digits: Decimal places for percentages.
        wrap_width: Character wrap width for first column in PNG output.
        include_test_stats_csv: If True, CSV includes inferential-test columns:
            statistic_type, test_statistic, p_value.
        continuous_test: Test for continuous variables — "anova" (default) or
            "mannwhitneyu" (2-group comparisons only).

    Returns:
        pd.DataFrame: Display table dataframe (first col 'Characteristic').
    """
    if sections is None:
        raise ValueError("`sections` is required.")

    sections = OrderedDict(sections)
    resolved_columns = _resolve_columns(dataframe, columns, include_total=include_total)

    row_df = _build_table_rows(
        dataframe=dataframe,
        resolved_columns=resolved_columns,
        sections=sections,
        continuous_digits=continuous_digits,
        pct_digits=pct_digits,
    )
    display_df = row_df[["Characteristic", *resolved_columns.keys()]].copy()
    csv_df = display_df.copy()

    row_df_with_stats = row_df
    if include_test_stats_csv:
        row_df_with_stats = _attach_row_level_test_stats(
            row_df=row_df,
            dataframe=dataframe,
            resolved_columns=resolved_columns,
            continuous_test=continuous_test,
            include_omnibus_chi_square=include_omnibus_chi_square,
                    )

        display_df["P-value"] = [
            _format_p(val, av)
            for val, av in zip(
                row_df_with_stats["chi_square_p_value"],
                row_df_with_stats["assumption_violated"],
            )
        ]

        # Add full test stats to CSV outputs
        csv_df["statistic_type"] = row_df_with_stats["statistic_type"]
        csv_df["test_statistic"] = row_df_with_stats["test_statistic"]
        csv_df["p_value"] = row_df_with_stats["p_value"]
        csv_df["assumption_violated"] = row_df_with_stats["assumption_violated"]
        csv_df["P-value"] = display_df["P-value"]

    styler = _build_styler(display_df, row_df_with_stats, table_caption)

    if savepath:
        base = Path(savepath)
        if base.suffix.lower() in {".csv", ".png", ".html", ".tex"}:
            base = base.with_suffix("")
        base.parent.mkdir(parents=True, exist_ok=True)

        formats = {fmt.lower().lstrip(".") for fmt in save_formats}
        if "csv" in formats:
            csv_df.to_csv(base.with_suffix(".csv"), index=False)
        if "html" in formats:
            base.with_suffix(".html").write_text(styler.to_html(), encoding="utf-8")
        if "tex" in formats:
            latex_df = display_df.replace({"%": r"\%"}, regex=True)
            latex_text = latex_df.to_latex(
                index=False,
                escape=False,
                caption=table_caption if table_caption else None,
                label=table_label if table_label else None,
                column_format="l" + ("c" * len(resolved_columns)),
            )
            base.with_suffix(".tex").write_text(latex_text, encoding="utf-8")
        if "png" in formats:
            _save_png_table(
                display_df=display_df,
                row_meta=row_df,
                savepath=base.with_suffix(".png"),
                table_caption=table_caption,
                wrap_width=wrap_width,
            )

    return display_df


def _build_styler_thickonly(display_df: pd.DataFrame, row_meta: pd.DataFrame, table_caption: str):
    """Like _build_styler but without per-row thin bottom borders on data rows."""

    def style_row(row: pd.Series):
        row_type = row_meta.loc[row.name, "row_type"]
        if row_type in {"section_header", "section_break"}:
            return [
                "font-style: italic; font-weight: 700; background-color: #f4f4f4; border-top: 1px solid #555;"
            ] * len(row)
        if row_type == "binary_indicator":
            return ["font-weight: 600;"] + [""] * (len(row) - 1)
        return [""] * len(row)

    table_styles = [
        {
            "selector": "table",
            "props": [
                ("border-collapse", "collapse"),
                ("border-top", "2px solid #222"),
                ("border-bottom", "2px solid #222"),
                ("width", "100%"),
            ],
        },
        {
            "selector": "th",
            "props": [
                ("font-family", "Times New Roman, serif"),
                ("font-size", "12pt"),
                ("font-weight", "700"),
                ("border-bottom", "1.5px solid #333"),
                ("padding", "6px 8px"),
                ("text-align", "center"),
                ("border-left", "none"),
                ("border-right", "none"),
            ],
        },
        {
            "selector": "td",
            "props": [
                ("font-family", "Times New Roman, serif"),
                ("font-size", "11pt"),
                ("padding", "4px 8px"),
                ("border-left", "none"),
                ("border-right", "none"),
            ],
        },
        {"selector": "tbody td:first-child", "props": [("text-align", "left")]},
    ]
    if "P-value" in display_df.columns:
        col_idx = list(display_df.columns).index("P-value")
        table_styles.append(
            {
                "selector": f"th.col_heading.level0.col{col_idx}",
                "props": [("font-style", "italic")],
            }
        )

    styler = (
        display_df.style.hide(axis="index")
        .set_table_styles(table_styles)
        .apply(style_row, axis=1)
        .set_properties(subset=display_df.columns[1:], **{"text-align": "center"})
    )
    if table_caption:
        styler = styler.set_caption(table_caption)
    return styler


def _save_png_table_thickonly(
    display_df: pd.DataFrame,
    row_meta: pd.DataFrame,
    savepath: Path,
    table_caption: str,
    wrap_width: int,
):
    """Like _save_png_table but regular data rows have no bottom border (thick lines only)."""
    plt.rcParams["font.family"] = "Arial"
    wrapped = display_df.copy()
    _first_col = wrapped.columns[0]
    wrapped[_first_col] = wrapped[_first_col].fillna("").map(
        lambda x: "\n".join(textwrap.wrap(str(x), width=wrap_width)) if str(x) else ""
    )

    n_rows = len(wrapped)
    n_cols = len(wrapped.columns)
    fig_width = max(3 + n_cols, 3.5 + (n_cols * 2.25))
    fig_height = max(3.8, 0.5 + (n_rows * 0.16))

    fig, ax = plt.subplots(figsize=(fig_width, fig_height))
    ax.axis("off")

    table = ax.table(
        cellText=wrapped.values,
        colLabels=wrapped.columns,
        loc="upper left",
        cellLoc="left",
        colLoc="center",
    )
    if "P-value" in wrapped.columns:
        col_idx = list(wrapped.columns).index("P-value")
        table[(0, col_idx)].get_text().set_fontstyle("italic")
    table.auto_set_font_size(False)
    fontsize = 10.5 if n_rows <= 35 else 9.5
    table.set_fontsize(fontsize)

    if n_cols <= 2:
        first_col_width = 0.86
    elif n_cols <= 4:
        first_col_width = 0.44
    else:
        first_col_width = 0.39
    remaining_width = 1.0 - first_col_width
    other_col_width = remaining_width / (n_cols - 1)

    row_weights = []
    for row_i in range(n_rows):
        line_count = max(1, wrapped.iloc[row_i, 0].count("\n") + 1)
        row_type = row_meta.iloc[row_i]["row_type"]
        weight = 1.0 + 0.78 * (line_count - 1)
        if row_type in {"section_header", "section_break"}:
            weight += 0.28
        row_weights.append(weight)

    header_weight = 1.35
    table_fill = 0.95
    per_weight = table_fill / (header_weight + sum(row_weights))
    header_height = header_weight * per_weight

    for (row_i, col_i), cell in table.get_celld().items():
        cell.set_facecolor("white")
        cell.set_edgecolor("#222222")
        cell.set_width(first_col_width if col_i == 0 else other_col_width)

        if row_i == 0:
            cell.visible_edges = "TB"
            cell.set_text_props(weight="bold", ha="center", va="center")
            cell.set_linewidth(1.5)
            cell.set_height(header_height)
            continue

        row_type = row_meta.iloc[row_i - 1]["row_type"]
        row_height = row_weights[row_i - 1] * per_weight
        cell.set_height(row_height)

        if col_i == 0:
            cell.set_text_props(ha="left")
        else:
            cell.set_text_props(ha="center")

        if row_type in {"section_header", "section_break"}:
            cell.visible_edges = "TB"
            if col_i == 0:
                cell.set_text_props(style="italic", weight="bold", ha="left")
            cell.set_linewidth(1.0)
        else:
            cell.visible_edges = ""

    fig.subplots_adjust(left=0.02, right=0.98, top=0.95, bottom=0.02)
    fig.savefig(savepath, dpi=400, bbox_inches="tight")
    plt.close(fig)


def generate_publication_table_thickonly(
    dataframe: pd.DataFrame,
    columns: dict | None = None,
    sections: OrderedDict | dict | None = None,
    table_caption: str = "",
    table_label: str = "tab:summary",
    savepath: str | None = None,
    save_formats=("png", "html", "tex", "csv"),
    include_total: bool = True,
    continuous_digits: int = 2,
    pct_digits: int = 1,
    wrap_width: int = 60,
    include_test_stats_csv: bool = True,
    continuous_test: str = "anova",
    include_omnibus_chi_square: bool = True,
    row_label_name: str = "Characteristic",
):
    """
    Like generate_publication_table but PNG/HTML output shows only thick structural
    lines (header top/bottom, section header top/bottom) — no thin per-row borders.

    All arguments are identical to generate_publication_table.
    """
    if sections is None:
        raise ValueError("`sections` is required.")

    sections = OrderedDict(sections)
    resolved_columns = _resolve_columns(dataframe, columns, include_total=include_total)

    row_df = _build_table_rows(
        dataframe=dataframe,
        resolved_columns=resolved_columns,
        sections=sections,
        continuous_digits=continuous_digits,
        pct_digits=pct_digits,
    )
    display_df = row_df[["Characteristic", *resolved_columns.keys()]].copy()
    if row_label_name != "Characteristic":
        display_df = display_df.rename(columns={"Characteristic": row_label_name})
    csv_df = display_df.copy()

    row_df_with_stats = row_df
    if include_test_stats_csv:
        row_df_with_stats = _attach_row_level_test_stats(
            row_df=row_df,
            dataframe=dataframe,
            resolved_columns=resolved_columns,
            continuous_test=continuous_test,
            include_omnibus_chi_square=include_omnibus_chi_square,
                    )

        display_df["P-value"] = [
            _format_p(val, av)
            for val, av in zip(
                row_df_with_stats["chi_square_p_value"],
                row_df_with_stats["assumption_violated"],
            )
        ]

        csv_df["statistic_type"] = row_df_with_stats["statistic_type"]
        csv_df["test_statistic"] = row_df_with_stats["test_statistic"]
        csv_df["p_value"] = row_df_with_stats["p_value"]
        csv_df["assumption_violated"] = row_df_with_stats["assumption_violated"]
        csv_df["P-value"] = display_df["P-value"]

    styler = _build_styler_thickonly(display_df, row_df_with_stats, table_caption)

    if savepath:
        base = Path(savepath)
        if base.suffix.lower() in {".csv", ".png", ".html", ".tex"}:
            base = base.with_suffix("")
        base.parent.mkdir(parents=True, exist_ok=True)

        formats = {fmt.lower().lstrip(".") for fmt in save_formats}
        if "csv" in formats:
            csv_df.to_csv(base.with_suffix(".csv"), index=False)
        if "html" in formats:
            base.with_suffix(".html").write_text(styler.to_html(), encoding="utf-8")
        if "tex" in formats:
            latex_df = display_df.replace({"%": r"\%"}, regex=True)
            latex_text = latex_df.to_latex(
                index=False,
                escape=False,
                caption=table_caption if table_caption else None,
                label=table_label if table_label else None,
                column_format="l" + ("c" * len(resolved_columns)),
            )
            base.with_suffix(".tex").write_text(latex_text, encoding="utf-8")
        if "png" in formats:
            _save_png_table_thickonly(
                display_df=display_df,
                row_meta=row_df,
                savepath=base.with_suffix(".png"),
                table_caption=table_caption,
                wrap_width=wrap_width,
            )
        if "docx" in formats:
            if not _HAS_DOCX:
                raise ImportError("python-docx is required for .docx output: pip install python-docx")
            _save_docx_table_thickonly(
                display_df=display_df,
                row_meta=row_df,
                savepath=base.with_suffix(".docx"),
            )

    return display_df


def _build_styler_thickonly_combined(
    display_df: pd.DataFrame,
    row_meta: pd.DataFrame,
    table_caption: str,
    separator_col_idx: int,
    col_display_labels: dict | None = None,
):
    """
    Like _build_styler_thickonly but adds a thick vertical right-border after
    `separator_col_idx` to separate the left and right column groups.

    `col_display_labels` maps internal column names to display names (e.g.
    ``{'__r_Total': 'Total', '__r_P-value': 'P-value'}``).  When provided,
    ``styler.relabel_index`` is used so that the rendered HTML shows the
    human-readable headers WITHOUT actually renaming the underlying DataFrame
    (which would create duplicate column names and trigger a Styler KeyError).
    """
    # Resolve position-indexed display names.  Using the display names here
    # (rather than the internal names) ensures the P-value italic rule fires
    # for both the left-group and the right-group P-value columns.
    if col_display_labels:
        display_col_names = [col_display_labels.get(c, c) for c in display_df.columns]
    else:
        display_col_names = list(display_df.columns)

    def style_row(row: pd.Series):
        row_type = row_meta.loc[row.name, "row_type"]
        if row_type in {"section_header", "section_break"}:
            return [
                "font-style: italic; font-weight: 700; background-color: #f4f4f4; border-top: 1px solid #555;"
            ] * len(row)
        if row_type == "binary_indicator":
            return ["font-weight: 600;"] + [""] * (len(row) - 1)
        return [""] * len(row)

    table_styles = [
        {
            "selector": "table",
            "props": [
                ("border-collapse", "collapse"),
                ("border-top", "2px solid #222"),
                ("border-bottom", "2px solid #222"),
                ("width", "100%"),
            ],
        },
        {
            "selector": "th",
            "props": [
                ("font-family", "Times New Roman, serif"),
                ("font-size", "12pt"),
                ("font-weight", "700"),
                ("border-bottom", "1.5px solid #333"),
                ("padding", "6px 8px"),
                ("text-align", "center"),
                ("border-left", "none"),
                ("border-right", "none"),
            ],
        },
        {
            "selector": "td",
            "props": [
                ("font-family", "Times New Roman, serif"),
                ("font-size", "11pt"),
                ("padding", "4px 8px"),
                ("border-left", "none"),
                ("border-right", "none"),
            ],
        },
        {"selector": "tbody td:first-child", "props": [("text-align", "left")]},
        # Thick vertical separator: right-border on the last left-group column
        {
            "selector": f"th.col_heading.level0.col{separator_col_idx}",
            "props": [("border-right", "2px solid #333")],
        },
        {
            "selector": f"td.col{separator_col_idx}",
            "props": [("border-right", "2px solid #333")],
        },
        # Thin right border on the Characteristic column to visually separate
        # row labels from the first data column.
        {
            "selector": "th.col_heading.level0.col0",
            "props": [("border-right", "1px solid #555")],
        },
        {
            "selector": "td.col0",
            "props": [("border-right", "1px solid #555")],
        },
    ]

    # Italicize every P-value header (check display names so the right-group
    # P-value column is caught even when its internal name is '__r_P-value').
    for col_idx, col_name in enumerate(display_col_names):
        if str(col_name).strip().startswith("P-value"):
            table_styles.append(
                {
                    "selector": f"th.col_heading.level0.col{col_idx}",
                    "props": [("font-style", "italic")],
                }
            )

    styler = (
        display_df.style.hide(axis="index")
        .set_table_styles(table_styles)
        .apply(style_row, axis=1)
        .set_properties(subset=display_df.columns[1:], **{"text-align": "center"})
    )
    # Set human-readable column headers without renaming the underlying DataFrame.
    # Renaming would produce duplicate column names (e.g. two "Total", two "P-value")
    # which causes Styler.apply to raise a KeyError on non-unique columns.
    if col_display_labels:
        styler = styler.relabel_index(display_col_names, axis=1)
    if table_caption:
        styler = styler.set_caption(table_caption)
    return styler


def _save_png_table_thickonly_combined(
    display_df: pd.DataFrame,
    row_meta: pd.DataFrame,
    savepath: Path,
    table_caption: str,
    wrap_width: int,
    separator_col_idx: int,
    col_display_labels: dict | None = None,
):
    """
    Like _save_png_table_thickonly but draws a thick vertical line on the right
    edge of `separator_col_idx` (0-indexed in the rendered table, where col 0 is
    Characteristic) to visually separate the two column groups.

    col_display_labels maps internal column names → display strings for headers.
    Duplicate display labels (e.g. two "Total" columns) are fine for rendering.
    """
    plt.rcParams["font.family"] = "Arial"
    wrapped = display_df.copy()
    wrapped["Characteristic"] = wrapped["Characteristic"].fillna("").map(
        lambda x: "\n".join(textwrap.wrap(str(x), width=wrap_width)) if str(x) else ""
    )

    # Apply display-label remapping for headers (allows duplicate display names)
    display_col_names = [
        col_display_labels.get(c, c) if col_display_labels else c
        for c in wrapped.columns
    ]

    n_rows = len(wrapped)
    n_cols = len(wrapped.columns)
    # Combined tables are wide: use compact per-column sizing
    fig_width = max(10, 3.5 + (n_cols * 1.9))

    # Identify rows immediately before / after each section boundary so we can
    # add targeted padding and pin their text away from the section lines.
    #   preceding_section : data rows whose NEXT row is a section_header/break
    #                        → extra height at bottom  → text pinned to visual top
    #   first_data_after  : data rows whose PRIOR row is a section_header/break
    #                        → extra height at top     → text pinned to visual bottom
    preceding_section = set()
    first_data_after  = set()
    for row_i in range(n_rows - 1):
        if row_meta.iloc[row_i + 1]["row_type"] in {"section_header", "section_break"}:
            preceding_section.add(row_i)
    for row_i in range(1, n_rows):
        if row_meta.iloc[row_i - 1]["row_type"] in {"section_header", "section_break"}:
            first_data_after.add(row_i)

    # Build row_weights with ~50px padding bonuses around section dividers.
    #   section_header/break : +1.4 bonus  → ~50px above AND below the header text
    #   rows adjacent to a section line : +0.7 bonus → ~50px of space at the boundary
    # (Exact pixel equivalence depends on total weight; these values are calibrated
    # empirically for the combined split tables at 400 DPI.)
    row_weights = []
    for row_i in range(n_rows):
        line_count = max(1, wrapped.iloc[row_i, 0].count("\n") + 1)
        row_type = row_meta.iloc[row_i]["row_type"]
        weight = 1.0 + 0.78 * (line_count - 1)
        if row_type in {"section_header", "section_break"}:
            weight += 0.7   # ~25px padding above and below the section header text
        if row_i in preceding_section:
            weight += 0.35  # ~25px of space at the bottom of this row before the line
        if row_i in first_data_after:
            weight += 0.35  # ~25px of space at the top of this row after the line
        row_weights.append(weight)

    # Grow the figure height to absorb the extra whitespace so data rows don't
    # shrink to the point of being hard to read.
    n_section_boundaries = len(preceding_section)
    fig_height = max(3.8, 0.5 + (n_rows * 0.16) + n_section_boundaries * 0.275)

    fig, ax = plt.subplots(figsize=(fig_width, fig_height))
    ax.axis("off")

    table = ax.table(
        cellText=wrapped.values,
        colLabels=display_col_names,
        loc="upper left",
        cellLoc="left",
        colLoc="center",
    )
    # Italicize every P-value header
    for col_idx, col_name in enumerate(display_col_names):
        if str(col_name).strip().startswith("P-value"):
            table[(0, col_idx)].get_text().set_fontstyle("italic")

    table.auto_set_font_size(False)
    fontsize = 10.5 if n_rows <= 35 else 9.5
    table.set_fontsize(fontsize)

    # Compact first column for wide combined tables
    first_col_width = 0.28
    remaining_width = 1.0 - first_col_width
    other_col_width = remaining_width / (n_cols - 1)

    header_weight = 1.35
    table_fill = 0.95
    per_weight = table_fill / (header_weight + sum(row_weights))
    header_height = header_weight * per_weight

    for (row_i, col_i), cell in table.get_celld().items():
        cell.set_facecolor("white")
        cell.set_edgecolor("#222222")
        cell.set_width(first_col_width if col_i == 0 else other_col_width)

        is_separator = (col_i == separator_col_idx)
        is_char_col  = (col_i == 0)   # thin right border separates Characteristic from data

        if row_i == 0:
            edges = "TBR" if (is_separator or is_char_col) else "TB"
            cell.visible_edges = edges
            cell.set_text_props(weight="bold", ha="center", va="center")
            cell.set_linewidth(1.5)
            cell.set_height(header_height)
            continue

        row_type = row_meta.iloc[row_i - 1]["row_type"]
        row_height = row_weights[row_i - 1] * per_weight
        cell.set_height(row_height)

        # All data rows use va='center'; the extra height in preceding/first_data_after
        # rows creates symmetric whitespace above and below the text, which still
        # produces a visible gap at each section boundary without pushing text to
        # cell edges (which caused apparent overlap in adjacent rows).
        base_ha = "left" if col_i == 0 else "center"
        if row_meta.iloc[row_i - 1]["row_type"] not in {"section_header", "section_break"}:
            cell.set_text_props(ha=base_ha, va="center")
        else:
            # section_header / section_break: bold italic
            if col_i == 0:
                cell.set_text_props(style="italic", weight="bold", ha="left", va="center")
            else:
                cell.set_text_props(ha="center", va="center")

        if row_type in {"section_header", "section_break"}:
            edges = "TBR" if (is_separator or is_char_col) else "TB"
            cell.visible_edges = edges
            cell.set_linewidth(1.5)
        else:
            edges = "R" if (is_separator or is_char_col) else ""
            cell.visible_edges = edges
            if is_separator:
                cell.set_linewidth(1.5)

    fig.subplots_adjust(left=0.02, right=0.98, top=0.95, bottom=0.02)
    fig.savefig(savepath, dpi=400, bbox_inches="tight")
    plt.close(fig)


def generate_combined_split_table_thickonly(
    dataframe: pd.DataFrame,
    columns_left: dict,
    columns_right: dict,
    sections,
    dataframe_right: pd.DataFrame | None = None,
    table_caption: str = "",
    table_label: str = "tab:summary",
    savepath: str | None = None,
    save_formats=("png", "html", "tex", "csv"),
    include_total: bool = True,
    continuous_digits: int = 2,
    pct_digits: int = 1,
    wrap_width: int = 40,
    continuous_test_left: str = "mannwhitneyu",
    continuous_test_right: str = "mannwhitneyu",
    include_omnibus_chi_square: bool = True,
):
    """
    Build a combined descriptive table with two side-by-side column groups
    separated by a thick vertical line in PNG/HTML output.

    Typical use: left group = HPPD split (PPA- vs PPA+), right group = CAPS
    split (CAPS- vs CAPS+).  The Characteristic column is shared.

    Args:
        dataframe:            Source dataframe for the LEFT column group.
        columns_left:         Dict of column-label → boolean mask for left group.
        columns_right:        Dict of column-label → boolean mask for right group.
        sections:             Same OrderedDict format as generate_publication_table_thickonly.
        dataframe_right:      Source dataframe for the RIGHT column group.  If None,
                              falls back to `dataframe`.  Pass a different df when the
                              two groups use different row-subsets (e.g. CAPS-valid only).
        table_caption:        Table caption string.
        table_label:          LaTeX label.
        savepath:             Base path without extension.
        save_formats:         Output format tuple.
        include_total:        Append "Total" column to each group.
        continuous_digits:    Decimal places for mean +/- SD.
        pct_digits:           Decimal places for percentages.
        wrap_width:           Char wrap for Characteristic column in PNG (default 40
                              — shorter than single-group tables because the first col
                              is narrower in wide combined layouts).
        continuous_test_left: Statistical test for the left group ("mannwhitneyu" or "anova").
        continuous_test_right: Statistical test for the right group.

    Returns:
        pd.DataFrame: Combined display dataframe with internal column names.
    """
    if sections is None:
        raise ValueError("`sections` is required.")
    if dataframe_right is None:
        dataframe_right = dataframe

    sections = OrderedDict(sections)

    resolved_left  = _resolve_columns(dataframe,       columns_left,  include_total=include_total)
    resolved_right = _resolve_columns(dataframe_right, columns_right, include_total=include_total)

    row_df_left  = _build_table_rows(dataframe,       resolved_left,  sections, continuous_digits, pct_digits)
    row_df_right = _build_table_rows(dataframe_right, resolved_right, sections, continuous_digits, pct_digits)

    # Compute p-values for each group independently
    row_stats_left  = _attach_row_level_test_stats(
        row_df_left,  dataframe,       resolved_left,  continuous_test_left,
        include_omnibus_chi_square=include_omnibus_chi_square,
            )
    row_stats_right = _attach_row_level_test_stats(
        row_df_right, dataframe_right, resolved_right, continuous_test_right,
        include_omnibus_chi_square=include_omnibus_chi_square,
            )

    # Prefix right-side column keys that collide with left-side keys (e.g. "Total",
    # later "P-value").  These internal keys are remapped to their original display
    # labels at render time via col_display_labels.
    left_data_keys = list(resolved_left.keys())

    def _right_key(k):
        return f"__r_{k}" if (k in left_data_keys or k == "P-value") else k

    right_internal_keys = [_right_key(k) for k in resolved_right.keys()]
    # map internal → display label for any renamed right-side key
    col_display_labels = {_right_key(k): k for k in resolved_right.keys() if _right_key(k) != k}
    # P-value columns
    _pval_left  = "P-value"
    _pval_right = "__r_P-value"
    col_display_labels[_pval_right] = "P-value"

    # Build row data for right-side cols with renamed keys
    rename_map = {k: _right_key(k) for k in resolved_right.keys() if _right_key(k) != k}
    if rename_map:
        row_df_right = row_df_right.rename(columns=rename_map)

    # Assemble display_df: Characteristic | left_data | P-value | right_data | P-value
    display_df = row_df_left[["Characteristic", *left_data_keys]].copy()
    display_df[_pval_left] = [
        _format_p(val, av)
        for val, av in zip(
            row_stats_left["chi_square_p_value"],
            row_stats_left["assumption_violated"],
        )
    ]
    for internal_key in right_internal_keys:
        display_df[internal_key] = row_df_right[internal_key].values
    display_df[_pval_right] = [
        _format_p(val, av)
        for val, av in zip(
            row_stats_right["chi_square_p_value"],
            row_stats_right["assumption_violated"],
        )
    ]

    # The vertical separator lives after the last left-group column (P-value of left group).
    # Columns: 0=Characteristic, 1..len(left_data_keys)=left data, then P-value at index
    # 1+len(left_data_keys), then right group.
    separator_col_idx = 1 + len(left_data_keys)  # index of "P-value" (left group)

    # Row metadata comes from the left group (same row structure for both)
    row_meta = row_df_left

    if savepath:
        base = Path(savepath)
        if base.suffix.lower() in {".csv", ".png", ".html", ".tex"}:
            base = base.with_suffix("")
        base.parent.mkdir(parents=True, exist_ok=True)

        formats = {fmt.lower().lstrip(".") for fmt in save_formats}

        if "csv" in formats:
            csv_df = display_df.rename(columns=col_display_labels)
            csv_df.to_csv(base.with_suffix(".csv"), index=False)

        if "html" in formats:
            # Pass the internal-named display_df directly — _build_styler_thickonly_combined
            # uses relabel_index to show human-readable headers without renaming, which
            # would create duplicate column names and crash Styler.apply.
            styler = _build_styler_thickonly_combined(
                display_df, row_meta, table_caption, separator_col_idx,
                col_display_labels=col_display_labels,
            )
            base.with_suffix(".html").write_text(styler.to_html(), encoding="utf-8")

        if "tex" in formats:
            tex_df = display_df.rename(columns=col_display_labels).replace({"%": r"\%"}, regex=True)
            n_data_cols = len(display_df.columns) - 1
            latex_text = tex_df.to_latex(
                index=False,
                escape=False,
                caption=table_caption if table_caption else None,
                label=table_label if table_label else None,
                column_format="l" + ("c" * n_data_cols),
            )
            base.with_suffix(".tex").write_text(latex_text, encoding="utf-8")

        if "png" in formats:
            _save_png_table_thickonly_combined(
                display_df=display_df,
                row_meta=row_meta,
                savepath=base.with_suffix(".png"),
                table_caption=table_caption,
                wrap_width=wrap_width,
                separator_col_idx=separator_col_idx,
                col_display_labels=col_display_labels,
            )
        if "docx" in formats:
            if not _HAS_DOCX:
                raise ImportError("python-docx is required for .docx output: pip install python-docx")
            _save_docx_table_thickonly_combined(
                display_df=display_df,
                row_meta=row_meta,
                savepath=base.with_suffix(".docx"),
                separator_col_idx=separator_col_idx,
                col_display_labels=col_display_labels,
            )

    return display_df


# ── docx helpers (booktabs-style Word tables) ───────────────────────────────

_DOCX_FONT = "Arial"
_DOCX_FONT_SIZE_HEADER = 9.5   # pt
_DOCX_FONT_SIZE_DATA   = 9     # pt
_DOCX_BORDER_COLOR = "222222"
_DOCX_THICK  = "12"   # eighths of a point → 1.5 pt
_DOCX_MEDIUM = "8"    # eighths of a point → 1.0 pt


def _docx_remove_table_borders(table) -> None:
    """Remove all default table-level borders so cells control their own."""
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = _parse_xml(f'<w:tblPr {_nsdecls("w")}/>')
        table._tbl.insert(0, tblPr)
    for old in tblPr.findall(_qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(_parse_xml(
        f'<w:tblBorders {_nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    ))


def _docx_set_cell_borders(cell, *, top=None, bottom=None, right=None) -> None:
    """Set top/bottom/right borders on a cell (pass size string or None)."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(_qn('w:tcBorders')):
        tcPr.remove(old)
    parts = []
    for edge, sz in [("top", top), ("bottom", bottom),
                     ("left", None), ("right", right)]:
        if sz is not None:
            parts.append(
                f'<w:{edge} w:val="single" w:sz="{sz}" '
                f'w:space="0" w:color="{_DOCX_BORDER_COLOR}"/>'
            )
        else:
            parts.append(
                f'<w:{edge} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            )
    tcPr.append(_parse_xml(
        f'<w:tcBorders {_nsdecls("w")}>{"".join(parts)}</w:tcBorders>'
    ))


def _docx_write_cell(cell, text, *, bold=False, italic=False,
                     align="center", font_size=None) -> None:
    """Write plain text to a cell with formatting."""
    if font_size is None:
        font_size = _Pt(_DOCX_FONT_SIZE_DATA)
    p = cell.paragraphs[0]
    p.alignment = {
        "left": _WD_ALIGN.LEFT,
        "center": _WD_ALIGN.CENTER,
        "right": _WD_ALIGN.RIGHT,
    }.get(align, _WD_ALIGN.CENTER)
    pf = p.paragraph_format
    pf.space_before = _Pt(1)
    pf.space_after = _Pt(1)
    text = str(text) if not pd.isna(text) else ""
    run = p.add_run(text)
    run.font.name = _DOCX_FONT
    run.font.size = font_size
    run.font.bold = bold
    run.font.italic = italic


def _save_docx_table_thickonly(
    display_df: pd.DataFrame,
    row_meta: pd.DataFrame,
    savepath: Path,
) -> None:
    """Save a single-group descriptive table as a Word .docx.

    Matches the PNG styling: booktabs borders, bold headers, bold-italic
    section headers, italic P-value header, no data-cell bolding.
    """
    doc = _Document()
    section = doc.sections[0]
    section.top_margin = _Cm(2.0)
    section.bottom_margin = _Cm(2.0)
    section.left_margin = _Cm(2.0)
    section.right_margin = _Cm(2.0)

    n_rows = len(display_df)
    n_cols = len(display_df.columns)

    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.autofit = True
    _docx_remove_table_borders(table)

    header_font = _Pt(_DOCX_FONT_SIZE_HEADER)

    # Column headers
    for col_i, col_name in enumerate(display_df.columns):
        cell = table.rows[0].cells[col_i]
        is_pval = str(col_name).strip().startswith("P-value")
        _docx_write_cell(cell, col_name, bold=True, italic=is_pval,
                         align="left" if col_i == 0 else "center",
                         font_size=header_font)
        _docx_set_cell_borders(cell, top=_DOCX_THICK, bottom=_DOCX_THICK)

    # Data / section rows
    for row_i in range(n_rows):
        rtype = row_meta.iloc[row_i]["row_type"]
        is_section = rtype in {"section_header", "section_break"}
        is_last = (row_i == n_rows - 1)

        for col_i in range(n_cols):
            cell = table.rows[row_i + 1].cells[col_i]
            text = str(display_df.iloc[row_i, col_i])

            if is_section:
                _docx_write_cell(cell, text if col_i == 0 else "",
                                 bold=True, italic=True, align="left")
                _docx_set_cell_borders(cell, top=_DOCX_MEDIUM,
                                       bottom=_DOCX_MEDIUM)
            else:
                _docx_write_cell(cell, text,
                                 align="left" if col_i == 0 else "center")
                if is_last:
                    _docx_set_cell_borders(cell, bottom=_DOCX_THICK)

    # Merge section-header rows across all columns
    for row_i in range(n_rows):
        rtype = row_meta.iloc[row_i]["row_type"]
        if rtype not in {"section_header", "section_break"}:
            continue
        row = table.rows[row_i + 1]
        merged = row.cells[0].merge(row.cells[n_cols - 1])
        while len(merged.paragraphs) > 1:
            p = merged.paragraphs[-1]
            p._element.getparent().remove(p._element)
        merged.paragraphs[0].clear()
        label_text = str(display_df.iloc[row_i, 0])
        _docx_write_cell(merged, label_text, bold=True, italic=True,
                         align="left")
        _docx_set_cell_borders(merged, top=_DOCX_MEDIUM, bottom=_DOCX_MEDIUM)

    doc.save(str(savepath))


def _save_docx_table_thickonly_combined(
    display_df: pd.DataFrame,
    row_meta: pd.DataFrame,
    savepath: Path,
    separator_col_idx: int,
    col_display_labels: dict | None = None,
) -> None:
    """Save a combined two-group descriptive table as a Word .docx.

    Like the single-group version but adds a thick vertical right-border
    after ``separator_col_idx`` to separate the left and right column groups.
    Column headers use display labels from ``col_display_labels``.
    """
    doc = _Document()
    section = doc.sections[0]
    # Combined tables are wide — use landscape
    section.orientation = _WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.top_margin = _Cm(2.0)
    section.bottom_margin = _Cm(2.0)
    section.left_margin = _Cm(1.5)
    section.right_margin = _Cm(1.5)

    n_rows = len(display_df)
    n_cols = len(display_df.columns)

    # Display column names (may differ from internal names for right-group)
    display_col_names = [
        col_display_labels.get(c, c) if col_display_labels else c
        for c in display_df.columns
    ]

    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.autofit = True
    _docx_remove_table_borders(table)

    header_font = _Pt(_DOCX_FONT_SIZE_HEADER)

    # Column headers
    for col_i, col_name in enumerate(display_col_names):
        cell = table.rows[0].cells[col_i]
        is_pval = str(col_name).strip().startswith("P-value")
        is_separator = (col_i == separator_col_idx)
        _docx_write_cell(cell, col_name, bold=True, italic=is_pval,
                         align="left" if col_i == 0 else "center",
                         font_size=header_font)
        _docx_set_cell_borders(cell, top=_DOCX_THICK, bottom=_DOCX_THICK,
                               right=_DOCX_THICK if is_separator else None)

    # Data / section rows
    for row_i in range(n_rows):
        rtype = row_meta.iloc[row_i]["row_type"]
        is_section = rtype in {"section_header", "section_break"}
        is_last = (row_i == n_rows - 1)

        for col_i in range(n_cols):
            cell = table.rows[row_i + 1].cells[col_i]
            is_separator = (col_i == separator_col_idx)

            if is_section:
                text = str(display_df.iloc[row_i, col_i])
                _docx_write_cell(cell, text if col_i == 0 else "",
                                 bold=True, italic=True, align="left")
                _docx_set_cell_borders(
                    cell, top=_DOCX_MEDIUM, bottom=_DOCX_MEDIUM,
                    right=_DOCX_THICK if is_separator else None,
                )
            else:
                text = str(display_df.iloc[row_i, col_i])
                _docx_write_cell(cell, text,
                                 align="left" if col_i == 0 else "center")
                bottom = _DOCX_THICK if is_last else None
                _docx_set_cell_borders(
                    cell, bottom=bottom,
                    right=_DOCX_THICK if is_separator else None,
                )

    # Merge section-header rows across all columns
    for row_i in range(n_rows):
        rtype = row_meta.iloc[row_i]["row_type"]
        if rtype not in {"section_header", "section_break"}:
            continue
        row = table.rows[row_i + 1]
        merged = row.cells[0].merge(row.cells[n_cols - 1])
        while len(merged.paragraphs) > 1:
            p = merged.paragraphs[-1]
            p._element.getparent().remove(p._element)
        merged.paragraphs[0].clear()
        label_text = str(display_df.iloc[row_i, 0])
        _docx_write_cell(merged, label_text, bold=True, italic=True,
                         align="left")
        _docx_set_cell_borders(merged, top=_DOCX_MEDIUM, bottom=_DOCX_MEDIUM)

    doc.save(str(savepath))


__all__ = [
    "generate_publication_table",
    "generate_publication_table_thickonly",
    "generate_combined_split_table_thickonly",
]
