"""
_docx_helper.py — shared Word-table builder for supplement tables.

Generates publication-style .docx tables matching the matplotlib PNG format:
  - Booktabs-style borders (thick rules on header top/bottom and table close;
    medium rules on section-header rows; no interior rules)
  - Bold column headers, bold-italic section headers
  - Amber cell shading for P ≥ threshold
  - Subscript rendering for special text (e.g., r_{rb})
  - Merged section-header rows
  - Arial font throughout

Google Docs imports .docx cleanly and preserves all of the above.
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


FONT_NAME = "Arial"
FONT_SIZE = Pt(9)
HEADER_FONT_SIZE = Pt(9.5)
AMBER_HEX = "FFF3CD"
BORDER_COLOR = "222222"

# Border sizes in eighths of a point: 12 ≈ 1.5pt, 8 = 1pt, 4 = 0.5pt
THICK = "12"
MEDIUM = "8"


# ── low-level XML helpers ────────────────────────────────────────────────────

def _remove_table_borders(table) -> None:
    """Remove all default table-level borders so cells control their own."""
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    ))


def _set_cell_borders(cell, *, top=None, bottom=None) -> None:
    """Set top/bottom borders on a cell.

    Pass ``top`` / ``bottom`` as a border-size string (e.g. THICK) or None.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    parts = []
    for edge, sz in [("top", top), ("bottom", bottom), ("left", None), ("right", None)]:
        if sz is not None:
            parts.append(
                f'<w:{edge} w:val="single" w:sz="{sz}" '
                f'w:space="0" w:color="{BORDER_COLOR}"/>'
            )
        else:
            parts.append(
                f'<w:{edge} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            )
    tcPr.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>{"".join(parts)}</w:tcBorders>'
    ))


def _set_cell_shading(cell, color_hex: str) -> None:
    """Set cell background fill colour."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    tcPr.append(parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    ))


# ── cell-text writer ────────────────────────────────────────────────────────

def _write_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    align: str = "center",
    font_size=FONT_SIZE,
    subscript_map: dict | None = None,
) -> None:
    """Write text to a cell, handling rich-text via *subscript_map*.

    *subscript_map* maps raw text strings to a list of segment dicts, each with
    keys ``text`` (str), and optional ``bold``, ``italic``, ``subscript`` (bool).
    """
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)

    text = str(text) if not pd.isna(text) else ""

    # Rich-text path (subscripts, mixed formatting)
    if subscript_map and text in subscript_map:
        for seg in subscript_map[text]:
            run = p.add_run(seg["text"])
            run.font.name = FONT_NAME
            run.font.size = font_size
            run.font.bold = seg.get("bold", bold)
            run.font.italic = seg.get("italic", italic)
            if seg.get("subscript"):
                run.font.subscript = True
        return

    # Plain-text path
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = font_size
    run.font.bold = bold
    run.font.italic = italic


# ── public API ───────────────────────────────────────────────────────────────

def save_docx_tables(
    panels: list[dict],
    output_path: Path | str,
    *,
    prob_threshold: float = 0.90,
    subscript_map: dict | None = None,
    landscape: bool = False,
) -> None:
    """Save one or more panel tables to a publication-style Word .docx.

    Parameters
    ----------
    panels : list of dict, each with keys:
        display_df   — pd.DataFrame of formatted display strings
        row_meta     — pd.DataFrame with a ``row_type`` column
                       ("section_header" or "data")
        prob_rows    — list[list] | None — per-cell probability floats for
                       amber highlighting (index 0-based into data columns)
        panel_label  — str | None — e.g. "a", "b"
    output_path : path to write the .docx file
    prob_threshold : cells with P ≥ this get amber shading + bold text
    subscript_map : maps raw text → rich-text segment list for subscript
                    rendering (e.g. ``{"$r_{rb}$": [...]}``)
    landscape : if True, set page orientation to landscape (for wide tables)
    """
    output_path = Path(output_path)
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        # Swap page width/height for landscape
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    else:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    for panel_idx, panel in enumerate(panels):
        display_df = panel["display_df"]
        row_meta   = panel["row_meta"]
        prob_rows  = panel.get("prob_rows")
        label      = panel.get("panel_label")

        # Gap between panels
        if panel_idx > 0:
            doc.add_paragraph()

        # Panel label
        if label:
            p = doc.add_paragraph()
            run = p.add_run(label)
            run.font.bold = True
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

        n_rows = len(display_df)
        n_cols = len(display_df.columns)

        # Create table: 1 header row + n_rows data/section rows
        table = doc.add_table(rows=n_rows + 1, cols=n_cols)
        table.autofit = True
        _remove_table_borders(table)

        # ── column-header row ────────────────────────────────────────────
        for col_i, col_name in enumerate(display_df.columns):
            cell = table.rows[0].cells[col_i]
            _write_cell(cell, col_name, bold=True,
                        align="left" if col_i == 0 else "center",
                        font_size=HEADER_FONT_SIZE,
                        subscript_map=subscript_map)
            _set_cell_borders(cell, top=THICK, bottom=THICK)

        # ── data / section rows ──────────────────────────────────────────
        for row_i in range(n_rows):
            rtype   = row_meta.iloc[row_i]["row_type"]
            is_last = (row_i == n_rows - 1)

            for col_i in range(n_cols):
                cell = table.rows[row_i + 1].cells[col_i]
                text = str(display_df.iloc[row_i, col_i])

                if rtype == "section_header":
                    # Will be overwritten after merge; write placeholder
                    _write_cell(cell, text if col_i == 0 else "",
                                bold=True, italic=True, align="left")
                    _set_cell_borders(cell, top=MEDIUM, bottom=MEDIUM)
                else:
                    # Check amber highlight
                    highlighted = False
                    if prob_rows is not None and col_i > 0:
                        p_idx = col_i - 1
                        probs = prob_rows[row_i]
                        if (probs is not None
                                and p_idx < len(probs)
                                and probs[p_idx] is not None
                                and not pd.isna(probs[p_idx])
                                and probs[p_idx] >= prob_threshold):
                            highlighted = True

                    _write_cell(cell, text, bold=highlighted,
                                align="left" if col_i == 0 else "center")

                    if is_last:
                        _set_cell_borders(cell, bottom=THICK)
                    # else: no borders (already cleared at table level)

        # ── merge section-header rows ────────────────────────────────────
        for row_i in range(n_rows):
            if row_meta.iloc[row_i]["row_type"] != "section_header":
                continue
            row = table.rows[row_i + 1]
            merged = row.cells[0].merge(row.cells[n_cols - 1])
            # Merging concatenates paragraphs; clear extras, keep first
            while len(merged.paragraphs) > 1:
                p = merged.paragraphs[-1]
                p._element.getparent().remove(p._element)
            # Rewrite text in the merged cell
            merged.paragraphs[0].clear()
            label_text = str(display_df.iloc[row_i, 0])
            _write_cell(merged, label_text, bold=True, italic=True,
                        align="left")
            _set_cell_borders(merged, top=MEDIUM, bottom=MEDIUM)

    doc.save(str(output_path))
